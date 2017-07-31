#! /usr/bin/env python3
# -*- coding: utf-8 -*-

from io import StringIO, BytesIO
import re

from docx import Document
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFResourceManager, process_pdf
from pyth.plugins.plaintext.writer import PlaintextWriter
from pyth.plugins.rtf15.reader import Rtf15Reader
import html2text
import olefile

import jieba
import jieba.analyse

from spider import download

def _parse_docx_table(table, text=''):
    for row in table.rows:
        for cell in row.cells:
            text += '\n\n'.join([
                paragraph.text for paragraph in cell.paragraphs
            ])

            for table in cell.tables:
                text += _parse_docx_table(table, text)

    return text

def readfile(file):
    try:
        if file.startswith('https://') or file.startswith('http://') or file.startswith('ftp://'):
            data = BytesIO(download(file))
        else:
            data = open(file, 'rb')

        if file.endswith('.caj') or file.endswith('.pdf'):
            with StringIO() as outfp:
                rsrcmgr = PDFResourceManager()
                device = TextConverter(rsrcmgr, outfp)
                process_pdf(rsrcmgr, device, data)
                return outfp.getvalue()
        elif file.endswith('.doc'):
            text = ''
            document = olefile.OleFileIO(data)

            wordDocument = document.openstream('WordDocument').read()

            # Parsing the WordDocument Stream
            # See https://msdn.microsoft.com/en-us/library/office/dd904907(v=office.14).aspx
            # And http://b2xtranslator.sourceforge.net/howtos/How_to_retrieve_text_from_a_binary_doc_file.pdf

            # Loading the FIB
            fib = wordDocument[:1472]

            # Loading and Parsing the piece table
            fcClx = int.from_bytes(fib[0x01A2:0x01A5], byteorder='little')
            lcbClx = int.from_bytes(fib[0x01A6:0x01A9], byteorder='little')

            tableFlag = ((int.from_bytes(fib[0x000A:0x000E], byteorder='little') & 0x0200) == 0x0200)
            tableName = ('0Table', '1Table')[tableFlag]

            table = document.openstream(tableName).read()

            clx = table[fcClx:fcClx+lcbClx]

            pos = 0
            pieceTable = ''
            lcbPieceTable = 0
            while True:
                if clx[pos] == 2:
                    # this entry is the piece table
                    lcbPieceTable = int.from_bytes(clx[pos+1:pos+5], byteorder='little')
                    pieceTable = clx[pos+5:pos+5+lcbPieceTable]
                    break
                elif clx[pos] == 1:
                    # skip this entry
                    pos = pos + 1 + 1 + ord(clx[pos + 1])
                else:
                    break

            i = 1
            pieceCount = (lcbPieceTable - 4) / 12
            while i <= pieceCount:
                cpStart = int.from_bytes(pieceTable[i*4:i*4+4], byteorder='little')
                cpEnd = int.from_bytes(pieceTable[(i+1)*4:(i+1)*4+4], byteorder='little')

                offsetPieceDescriptor = int(((pieceCount + 1) * 4) + (i * 8))
                pieceDescriptor = pieceTable[offsetPieceDescriptor:offsetPieceDescriptor+8]

                fcValue = int.from_bytes(pieceDescriptor[2:6], byteorder='little')
                isANSII = (fcValue & 0x40000000) == 0x40000000
                fc = fcValue & 0xBFFFFFFF

                encoding = ('utf-16', 'cp1252')[isANSII]
                cb = cpEnd - cpStart
                cb = (cb * 2, cb)[isANSII]
                text += wordDocument[fc:fc+cb].decode(encoding)

                i += 1

            return text
        elif file.endswith('.docx'):
            text = ''
            document = Document(data)

            text += '\n\n'.join([
                paragraph.text for paragraph in document.paragraphs
            ])

            for table in document.tables:
                text += _parse_docx_table(table, text)

            return text
        elif file.endswith('.htm') or file.endswith('.html'):
            html = html2text.HTML2Text()
            html.ignore_links = True
            return html.handle(data.read().decode('utf-8'))
        elif file.endswith('.rtf'):
            with BytesIO() as outfp:
                document = Rtf15Reader.read(data)
                return PlaintextWriter.write(document, outfp).getvalue()
        elif file.endswith('.txt'):
            return data.read()
        else:
            raise Exception('Unknown file extension')
    except:
        pass

def getkeywords(text):
    try:
        # TF-IDF
        return jieba.analyse.extract_tags(text, topK=10)
        # TextRank
        #return jieba.analyse.textrank(text, topK=10)
    except:
        pass

def getsegments(text):
    try:
        segments = jieba.lcut(text, cut_all=False, HMM=False)

        remove = u'\s!"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~＂＃＄％＆＇（）＊＋，－／：；＜＝＞＠［＼］＾＿｀｛｜｝～｟｠｢｣､、〃》「」『』【】〔〕〖〗〘〙〚〛〜〝〞〟〰〾〿–—‘’‛“”„‟…‧﹏'
        pattern = re.compile(r"[{}]".format(remove))

        return [s for s in segments if not pattern.match(s)]
    except:
        pass
